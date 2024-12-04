from re import X
import time
import io
from tkinter import Y
import openpyxl
from docx import Document
from selenium import webdriver
from selenium.webdriver.common.by import By
from docx.shared import Inches 
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait, Select
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.edge.service import Service
from selenium.webdriver.edge.options import Options
import mss
import numpy as np
import cv2
import pyautogui
from PIL import ImageGrab

########################################################################################################################################
screenshot_counter = 0
document = Document()

def ejecutar_acciones(archivo_excel):
    chrome_options = Options()
    chrome_options.add_argument("--inprivate")
    driver = webdriver.Edge(options=chrome_options)
    driver.maximize_window()

    def take_screenshot(step_name):
        global screenshot_counter
        screenshot_counter += 1
        time.sleep(3)

        with mss.mss() as sct:
            monitor = sct.monitors[1]
            screenshot = sct.grab(monitor)

            img_stream = io.BytesIO()
            img_array = np.array(screenshot)
            img_array = cv2.cvtColor(img_array, cv2.COLOR_BGRA2BGR)
            _, buffer = cv2.imencode('.png', img_array)
            img_stream.write(buffer)
            img_stream.seek(0)

            document.add_paragraph(f"{step_name}")
            picture = document.add_picture(img_stream)
            picture.width = Inches(6)  # Ancho ajustado a 5 pulgadas
            picture.height = Inches(3)
            

    wb = openpyxl.load_workbook(archivo_excel)
    sheet = wb.active  
    print("Datos cargados")

    def wait_for_element(by, value, timeout=35):
        return WebDriverWait(driver, timeout).until(EC.visibility_of_element_located((by, value)))
    
    # Leer acciones desde el Excel
    for row in sheet.iter_rows(min_row=2, values_only=True):
        # Imprimir el contenido de la fila
        print(f"Fila leída: {row}")

        # Asegúrate de que la fila tenga al menos 5 elementos
        if len(row) < 5:
            print(f"Fila incompleta: {row}. Se omitirá.")
            continue

        Accion, Elemento, ValorE, Opciones, Nombre = row

        if Accion == 'navigate':
            driver.get(ValorE)
            print("Iniciando....")
            take_screenshot("N/A")
        
        elif Accion == 'wait':
            wait_for_element(By.NAME, Elemento, int(Opciones))
            print("Espera....")
            take_screenshot(Nombre if Nombre else "N/A")
         
        elif Accion == 'waitxp':
            wait_for_element(By.XPATH, Elemento, int(Opciones))
            print("Espera....")
            take_screenshot(Nombre if Nombre else "N/A")
           
        elif Accion == 'write':
            wait_for_element(By.NAME, Elemento)  
            driver.find_element(By.NAME, Elemento).send_keys(ValorE)
            print("Escribiendo.....")
            take_screenshot(Nombre if Nombre else "N/A")

        elif Accion == 'writexp':
            wait_for_element(By.XPATH, Elemento)  
            driver.find_element(By.XPATH, Elemento).send_keys(ValorE)
            print("Escribiendo.....")
            take_screenshot(Nombre if Nombre else "N/A")
            
        elif Accion == 'click':
            wait_for_element(By.ID, Elemento).click()
            print("Dando Click....")
            take_screenshot(Nombre if Nombre else "N/A")

        elif Accion == 'clickxp':
            element = wait_for_element(By.XPATH, Elemento)
            driver.execute_script("window.scrollTo(0, arguments[0].offsetTop);", element)
            WebDriverWait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, Elemento)))
            element.click()
            print("Dando Click....")
            take_screenshot(Nombre if Nombre else "N/A")
            
        elif Accion == 'select':
            select_element = driver.find_element(By.XPATH, Elemento)
            select = Select(select_element)
            select.select_by_visible_text(ValorE)
            print("Seleccionando...")
            take_screenshot(Nombre if Nombre else "N/A")

        elif Accion == 'iframe':
            driver.switch_to.frame(0)
            print("Seleccionando iframe...")    
            take_screenshot(Nombre if Nombre else "N/A")

        elif Accion == 'iframeQ':
            driver.switch_to.default_content()
            print("Quitar iframe...")   
            take_screenshot("N/A")

        elif Accion == 'imagen':
            x, y = Elemento, ValorE
            pyautogui.moveTo(x, y, duration=1)
            pyautogui.click()   # Mover el mouse a (x, y) en 0.5 segundos
            print("click en la coordenada.")
            take_screenshot(Nombre if Nombre else "N/A")


        elif Accion == 'imagen2':
            x, y = Elemento, ValorE
            pyautogui.moveTo(x, y, duration=1)
            pyautogui.click()   # Mover el mouse a (x, y) en 0.5 segundos
            print("click en la coordenada.")
            take_screenshot(Nombre if Nombre else "N/A")


        elif Accion == 'imagen3':
            x, y = Elemento, ValorE
            pyautogui.moveTo(x, y, duration=1)
            pyautogui.click()   # Mover el mouse a (x, y) en 0.5 segundos
            print("click en la coordenada.")
            take_screenshot(Nombre if Nombre else "N/A")     

        try:
              document.save(r'C:\Users\gelver.beltran\Documents\pruebas\prueba.docx')   
        except Exception as e:
            print(f"Error al guardar el documento: {str(e)}")
  

archivo_excel = r"C:/Users/gelver.beltran/OneDrive - PEOPLE TECH LATIN S.A.S/Documentos/Python/Prospectoohpay.xlsx"
ejecutar_acciones(archivo_excel)
time.sleep(15)

